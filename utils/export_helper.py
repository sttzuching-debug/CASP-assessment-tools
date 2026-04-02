from docx import Document
from docx.shared import RGBColor
import io
from datetime import datetime

def create_casp_word_report(result_json, local_context="無", paper_name="未知文獻"):
    doc = Document()
    doc.add_heading('CASP 系統性回顧 (SRMA) 評讀報告', 0)
    doc.add_paragraph(f"評讀時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}\n評讀文獻：{paper_name}")
    doc.add_heading('📍 應用在地情境', level=2)
    doc.add_paragraph(local_context)
    
    for key, val in result_json.items():
        verdict = val.get('Verdict', "Can't tell")
        heading = doc.add_heading(f"{key}: {val.get('Question', '')} 【{verdict}】", level=2)
        doc.add_paragraph(f"💡 解析：{val.get('Reasoning_TW', '')}")
        if 'Auditor_Note' in val:
            doc.add_paragraph(f"🛡️ 稽核備註：{val['Auditor_Note']}")
            
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
